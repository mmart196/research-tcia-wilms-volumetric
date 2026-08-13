"""Render draft/manuscript.md from manuscript_template.md + data/results.json.

The template contains no results numbers — every statistic is a {token} filled
from results.json (itself written by analysis.py). This is the spec-R4 contract:
the manuscript literally cannot state a number the pipeline did not produce.

    python render_draft.py        # after analysis.py has run
"""
import json
import math
import os
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "data"

# Pediatric Radiology's technical check reported this DOI unresolvable on 2026-08-03.
# It was derived from the version DOI ...440 by subtraction, on the assumption that a
# Zenodo concept DOI is the version DOI minus one. Zenodo mints the concept DOI
# separately, so that assumption was wrong and the journal proved it empirically.
# Rendering is blocked until a human reads the real concept DOI off the Zenodo record
# page, because a broken data-availability link has already cost one technical-check
# round and would cost another. Set ALLOW_BAD_DOI=1 to build a draft package anyway;
# nothing built that way may be submitted.
KNOWN_BAD_DOI = "10.5281/zenodo.21608439"


def fmt(x, nd=1):
    return f"{x:.{nd}f}"


_ONES = ("zero", "one", "two", "three", "four", "five", "six", "seven", "eight", "nine")


def word(n):
    """Cureus number style: spell out 1-9, leave 10 and above as digits.

    Counts like "1 subject" have to read "one subject" to pass Cureus's
    formatting check, and a submission that fails it is deferred for the paid
    Preferred Editing service. The value still comes from results.json -- only
    its rendering changes -- so the provenance contract holds. The threshold is
    applied here rather than hand-typed into the template because these counts
    are data-derived and could cross 9 on a re-run.
    """
    n = int(n)
    return _ONES[n] if 0 <= n <= 9 else str(n)


def fmt_p(p):
    if p < 0.001:
        exp = math.floor(math.log10(p))
        mant = p / 10 ** exp
        return f"{mant:.1f} × 10<sup>{exp}</sup>"
    return f"{p:.3f}"


def main():
    res = json.load((DATA / "results.json").open(encoding="utf-8"))
    refs = json.load((ROOT / "draft" / "references.json").open(encoding="utf-8"))

    flow, coh, prim, subs = res["flow"], res["cohort"], res["primary"], res["subgroups_by_collection"]

    t = {
        # flow
        "F_META_ANY": flow["metadata_subjects_any_kidney_seg"],
        "F_META_PAIRED": flow["metadata_paired_pre_post"],
        "F_EXTRACTED": flow["extracted_subjects_any_usable_ct_volume"],
        "F_PAIRED": flow["paired_analyzed"],
        "F_EXCLUDED": flow["excluded_no_pre_or_post_volume"],
        # cohort
        "N": coh["n"],
        "AGE_MED": fmt(coh["age_median"]),
        "AGE_Q1": fmt(coh["age_iqr"][0]),
        "AGE_Q3": fmt(coh["age_iqr"][1]),
        "AGE_MIN": fmt(coh["age_min"]),
        "AGE_MAX": fmt(coh["age_max"]),
        "SEX_M": coh["sex_counts"].get("M", 0),
        "SEX_F": coh["sex_counts"].get("F", 0),
        "SEX_O": coh["sex_counts"].get("O", 0),
        "LAT_L": coh["laterality_counts"].get("left", 0),
        "LAT_R": coh["laterality_counts"].get("right", 0),
        "LAT_B": coh["laterality_counts"].get("bilateral", 0),
        "LAT_U": coh["laterality_counts"].get("unknown", 0),
        "C_0532": coh["collection_counts"].get("AREN0532", 0),
        "C_0533": coh["collection_counts"].get("AREN0533", 0),
        "C_0534": coh["collection_counts"].get("AREN0534", 0),
        # primary
        "PRE_MED": fmt(prim["pre_median_ml"]),
        "PRE_Q1": fmt(prim["pre_iqr_ml"][0]),
        "PRE_Q3": fmt(prim["pre_iqr_ml"][1]),
        "POST_MED": fmt(prim["post_median_ml"]),
        "POST_Q1": fmt(prim["post_iqr_ml"][0]),
        "POST_Q3": fmt(prim["post_iqr_ml"][1]),
        "PCT_MED": fmt(prim["pct_change_median"]),
        "PCT_Q1": fmt(prim["pct_change_iqr"][0]),
        "PCT_Q3": fmt(prim["pct_change_iqr"][1]),
        "N_DEC": prim["n_decreased"],
        "N_INC": prim["n_increased"],
        "W_STAT": fmt(prim["wilcoxon_statistic"], 0),
        "W_P": fmt_p(prim["wilcoxon_p"]),
        # subgroups
        "S0532_N": subs.get("AREN0532", {}).get("n", 0),
        "S0532_MED": fmt(subs.get("AREN0532", {}).get("median_pct_change", float("nan"))),
        "S0533_N": subs.get("AREN0533", {}).get("n", 0),
        "S0533_MED": fmt(subs.get("AREN0533", {}).get("median_pct_change", float("nan"))),
        "S0534_N": subs.get("AREN0534", {}).get("n", 0),
        "S0534_MED": fmt(subs.get("AREN0534", {}).get("median_pct_change", float("nan"))),
        # v2 additions
        "HL_EST": fmt(prim["hl_pseudomedian"]),
        "HL_LO": fmt(prim["hl_ci_95"][0]),
        "HL_HI": fmt(prim["hl_ci_95"][1]),
        "BILAT_SHARE": round(100 * subs.get("AREN0534", {}).get("n", 0) / max(coh["n"], 1)),
        "TABLE1_MD": res["table1_md"],
        "VVOL_N": res["volume_validation"]["n_rois"],
        "VVOL_RHO": fmt(res["volume_validation"].get("spearman_rho", float("nan")), 3),
        "VVOL_MEDDIFF": fmt(res["volume_validation"].get("median_abs_pct_diff", float("nan"))),
        "VVOL_PRE_SIGNED": fmt(res["volume_validation"].get("pre_signed_median_pct", float("nan"))),
        "VVOL_POST_SIGNED": fmt(res["volume_validation"].get("post_signed_median_pct", float("nan"))),
        "VVOL_PRE_ABS": fmt(res["volume_validation"].get("pre_abs_median_pct", float("nan"))),
        "VVOL_POST_ABS": fmt(res["volume_validation"].get("post_abs_median_pct", float("nan"))),
        "F_POSTOP": res["flow_postop_only_excluded"],
        # Cureus number style: spelled-out variants for counts that read as prose.
        "SEX_O_W": word(coh["sex_counts"].get("O", 0)),
        "LAT_U_W": word(coh["laterality_counts"].get("unknown", 0)),
        "F_POSTOP_W": word(res["flow_postop_only_excluded"]),
    }

    # Vancouver-style numbered reference list (all records PubMed-resolved)
    ref_lines = []
    for i, r in enumerate(refs, 1):
        auth = ", ".join(r["authors"][:6]) + (" et al" if len(r["authors"]) > 6 else "")
        vi = f";{r['volume']}" if r["volume"] else ""
        if r["issue"]:
            vi += f"({r['issue']})"
        pg = f":{r['pages']}" if r["pages"] else ""
        doi = f" doi:{r['doi']}" if r["doi"] else ""
        ref_lines.append(f"{i}. {auth}. {r['title']} {r['journal']}. {r['year']}{vi}{pg}.{doi} "
                         f"PMID: {r['pmid']}")
    t["REFERENCES"] = "\n".join(ref_lines)

    tpl = (ROOT / "draft" / "manuscript_template.md").read_text(encoding="utf-8")
    if KNOWN_BAD_DOI in tpl and not os.environ.get("ALLOW_BAD_DOI"):
        raise SystemExit(
            f"refusing to render: the template still cites {KNOWN_BAD_DOI}, which "
            "Pediatric Radiology reported unresolvable on 2026-08-03. Read the concept "
            "DOI off the Zenodo record page ('Cite all versions') and replace it. Do not "
            "derive it from the version DOI, and do not revert to ...440 -- that record "
            "still contains internal planning material. To build an unsubmittable draft "
            "package anyway: ALLOW_BAD_DOI=1 python code/render_draft.py"
        )
    out = tpl
    missing = []
    for k, v in t.items():
        tag = "{" + k + "}"
        if tag not in out and k != "REFERENCES" or (k == "REFERENCES" and "{REFERENCES}" not in out):
            missing.append(k)
        out = out.replace(tag, str(v))
    leftover = [tok for tok in out.split("{")[1:] if "}" in tok and tok.split("}")[0].isupper()]
    if leftover:
        raise SystemExit(f"template tokens left unfilled: {[t.split('}')[0] for t in leftover]}")
    if missing:
        print(f"note: tokens defined but not used in template: {missing}")

    dest = ROOT / "draft" / "manuscript.md"
    dest.write_text(out, encoding="utf-8")
    print(f"Wrote {dest} ({len(out)} chars)")

    # ---- JIIM double-blind version ----
    # Blind on the actual identifying TOKENS (name, email, institutions, GitHub
    # handle) with individual regexes, not one brittle multi-line literal block.
    # The previous version matched the whole author block as one exact string
    # including "MD — michael13414@gmail.com" (em dash); when the template's
    # separator drifted to "MD, michael13414@gmail.com" (comma) the .replace()
    # silently no-opped and the blinded manuscript shipped with full author
    # names, affiliations, and email -- caught only because a human happened to
    # read the file, not by anything in this pipeline. Same failure class as the
    # Zenodo-DOI leak fixed above; both are now assert-checked rather than
    # trusted to fire silently.
    blind = out
    author_block_re = re.compile(
        r"Rachel Velasco.*?michael13414@gmail\.com", re.DOTALL)
    blind, n_author = author_block_re.subn(
        "[Author information blinded for review]", blind, count=1)
    if n_author != 1:
        raise RuntimeError(
            "author-block blinding regex matched "
            f"{n_author} times (expected exactly 1) -- template text likely "
            "changed; the blinded manuscript would ship with author identity "
            "intact")
    blind = blind.replace("Velasco:", "[blinded]:").replace("Martinez:", "[blinded]:")
    blind = blind.replace("https://github.com/mmart196/research-tcia-wilms-volumetric",
                          "[repository link anonymized for review]")
    for leak in ("Rachel Velasco", "Michael Martinez", "michael13414@gmail.com",
                 "Washington University of Health and Science",
                 "TidalHealth Peninsula Regional",
                 "github.com/mmart196"):
        if leak in blind:
            raise RuntimeError(
                f"blinded manuscript still contains identifying text: {leak!r}")
    # Anonymize ANY Zenodo DOI, not one hard-coded number. The previous version
    # matched only 10.5281/zenodo.21608440; when the manuscript's DOI was changed
    # during the Pediatric Radiology retarget the replacement silently stopped
    # firing, and the blinded manuscript shipped a live DOI that resolves to a
    # record carrying the authors' names. Matching the pattern means a future DOI
    # change cannot reopen that hole.
    blind, n_doi = re.subn(r"10\.5281/zenodo\.\d+", "[DOI anonymized for review]", blind)
    if re.search(r"10\.5281/zenodo\.\d+", blind):
        raise RuntimeError("blinded manuscript still carries a resolvable Zenodo DOI")
    bdest = ROOT / "draft" / "manuscript_blinded.md"
    bdest.write_text(blind, encoding="utf-8")
    print(f"Wrote {bdest} (blinded)")

    # ---- .docx for both (python-docx, simple markdown subset) ----
    for src, name in ((out, "manuscript"), (blind, "manuscript_blinded")):
        write_docx(src, ROOT / "draft" / f"{name}.docx")
        print(f"Wrote {name}.docx")

    # ---- title page ----
    # Optional. The public reproducibility package ships the pipeline and the
    # manuscript, not the submission paperwork, so title-page.md is absent there
    # and this step used to crash anyone who cloned the repo and ran the renderer
    # -- the exact "every number regenerates from public code" claim the package
    # exists to support.
    tp_path = ROOT / "draft" / "title-page.md"
    if tp_path.exists():
        tp_src = _strip_frontmatter(tp_path.read_text(encoding="utf-8"))
        write_docx(tp_src, ROOT / "draft" / "title-page.docx")
        print("Wrote title-page.docx")
    else:
        print("note: no draft/title-page.md, skipping title page (expected in the public package)")


def _strip_frontmatter(text):
    """Drop a leading Pandoc-style '---\\n...\\n---' YAML block, if present.

    write_docx() has no notion of document metadata, so a leaked frontmatter
    block was rendered as three literal body paragraphs ("---", 'title: "..."',
    "---") at the top of every title page. Metadata belongs in the file, not
    in the reader's view of it.
    """
    lines = text.splitlines(keepends=True)
    if lines and lines[0].strip() == "---":
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                return "".join(lines[i + 1:]).lstrip("\n")
    return text


import re as _re_mod
_LIST_RE = _re_mod.compile(r"^\\d+[.)]\\s")


def _md_runs(par, text):
    """Add text to a paragraph, honouring **bold** spans."""
    import re as _re
    for i, seg in enumerate(_re.split(r"\*\*(.+?)\*\*", text)):
        if not seg:
            continue
        run = par.add_run(seg)
        if i % 2 == 1:
            run.bold = True


def _split_row(line):
    """A markdown pipe row -> list of cell strings."""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _is_sep(line):
    body = line.strip().strip("|").replace("|", "")
    return body and set(body.replace(" ", "")) <= set("-:")


def _add_table(doc, rows):
    """Real Word table from markdown pipe rows. Header bold, grid borders."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt
    body = [r for r in rows if not _is_sep(r)]
    if not body:
        return
    grid = [_split_row(r) for r in body]
    ncol = max(len(r) for r in grid)
    grid = [r + [""] * (ncol - len(r)) for r in grid]
    t = doc.add_table(rows=len(grid), cols=ncol)
    try:
        t.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for ri, row in enumerate(grid):
        for ci, cell in enumerate(row):
            c = t.cell(ri, ci)
            c.text = ""
            par = c.paragraphs[0]
            _md_runs(par, cell)
            for run in par.runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
    doc.add_paragraph()


def write_docx(md_text, path):
    """Markdown subset -> .docx. Pipe tables become REAL Word tables.

    Fixed 2026-07-31: this used to emit `doc.add_paragraph(line)` for every pipe
    row, so Table 1 reached the submission portal as raw markdown -
    "| Characteristic | AREN0532 (n=24) | ..." - visible in the uploaded file.
    """
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    # HTML comments are markdown metadata, not content - strip them before rendering.
    md_text = _re_mod.sub(r"<!--.*?-->", "", md_text, flags=_re_mod.S)
    lines = md_text.splitlines()
    i = 0
    buf = []

    def flush():
        # Markdown rule: consecutive non-blank lines are ONE paragraph. Emitting a
        # paragraph per source line is what made the cover letter break mid-sentence
        # at every 90-character wrap point. Fixed 2026-07-31.
        # Markdown rule: consecutive non-blank lines are ONE paragraph, joined by a
        # space. A line ending in two spaces is a HARD line break (signature and
        # address blocks need this) and becomes a Word line break, not a new paragraph.
        if not buf:
            return
        par = doc.add_paragraph()
        for k, raw in enumerate(buf):
            _md_runs(par, raw.strip())
            if k < len(buf) - 1:
                if raw.rstrip("\n").endswith("  ") and par.runs:
                    par.runs[-1].add_break()
                else:
                    par.add_run(" ")
        del buf[:]

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("|"):
            flush()
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            _add_table(doc, block)
            continue
        if line.startswith("### "):
            flush()
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            flush()
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            flush()
            doc.add_heading(line[2:], level=1)
        elif stripped.startswith("---"):
            flush()
        elif stripped.startswith(("- ", "* ")) or _LIST_RE.match(stripped):
            flush()
            _md_runs(doc.add_paragraph(style="List Paragraph"), stripped)
        elif not stripped:
            flush()
        else:
            buf.append(line)
        i += 1
    flush()
    doc.save(str(path))


if __name__ == "__main__":
    main()
